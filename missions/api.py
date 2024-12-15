from rest_framework.viewsets import GenericViewSet
from rest_framework import mixins, viewsets, serializers
from rest_framework.permissions import IsAuthenticated

from rest_framework.views import APIView
from rest_framework.response import Response

from django.contrib.auth.models import User

from django.db.models import Avg, Count, Min, Max
from rest_framework.decorators import action

from rest_framework.permissions import AllowAny
from django.contrib.auth import authenticate, login
from rest_framework import status

import pyotp

from django.contrib.auth import logout
from rest_framework.permissions import BasePermission
from django.utils import timezone

from django.core.cache import cache

from django.views.decorators.csrf import csrf_exempt

from openpyxl import Workbook
from docx import Document
from rest_framework.decorators import action
from django.http import HttpResponse

from missions.models import Mission, SpaceProgram, MissionType, LaunchSite, SpaceCraft
from missions.serializers import (
    MissionSerializer, 
    SpaceProgramSerializer, 
    MissionTypeSerializer, 
    LaunchSiteSerializer, 
    SpaceCraftSerializer,
    UserSerializer,
)

# Вьюсет для миссий
# какие статистические данные можно выводить на странице с миссиями?
class MissionsViewset(mixins.CreateModelMixin,
                      mixins.DestroyModelMixin,
                      mixins.UpdateModelMixin,
                      mixins.RetrieveModelMixin,
                      mixins.ListModelMixin,
                      GenericViewSet):
    queryset = Mission.objects.all()
    serializer_class = MissionSerializer

    def get_queryset(self):
        qs = super().get_queryset()

        if self.request.user.is_superuser:
            user_id = self.request.query_params.get('user_id')
            if user_id:
                qs = qs.filter(user_id=user_id)

        elif self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)

        else:
            return qs.none()

        name = self.request.query_params.get('name')
        if name:
            qs = qs.filter(name__icontains=name)

        date = self.request.query_params.get('date')
        if date:
            qs = qs.filter(launch_date=date)

        outcome = self.request.query_params.get('outcome')
        if outcome:
            qs = qs.filter(outcome__icontains=outcome)

        description = self.request.query_params.get('description')
        if description:
            qs = qs.filter(description__icontains=description)

        program = self.request.query_params.get('program')
        if program:
            qs = qs.filter(space_program_id=program)

        mission_type = self.request.query_params.get('type')
        if mission_type:
            qs = qs.filter(mission_type_id=mission_type)

        return qs

    @action(detail=False, methods=["GET"], url_path="export-excel")
    def export_to_excel(self, request):
        missions = self.get_queryset()
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Missions"

        # Заголовки
        sheet.append(["ID", "Название", "Дата", "Описание", "Программа", "Тип", "Исход"])

        # Данные
        for mission in missions:
            sheet.append([
                mission.id,
                mission.name,
                mission.launch_date,
                mission.description,
                mission.space_program.name if mission.space_program else "Нет программы",
                mission.mission_type.name if mission.mission_type else "Нет типа",
                mission.outcome if mission.outcome else "Нет данных"  # Добавляем исход
            ])

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response["Content-Disposition"] = 'attachment; filename="missions.xlsx"'
        workbook.save(response)
        return response

    @action(detail=False, methods=["GET"], url_path="export-word")
    def export_to_word(self, request):
        missions = self.get_queryset()
        document = Document()

        document.add_heading("Миссии", level=1)

        for mission in missions:
            document.add_heading(f"Миссия: {mission.name}", level=2)
            document.add_paragraph(f"Дата: {mission.launch_date}")
            document.add_paragraph(f"Описание: {mission.description}")
            document.add_paragraph(f"Программа: {mission.space_program.name if mission.space_program else 'Нет программы'}")
            document.add_paragraph(f"Тип: {mission.mission_type.name if mission.mission_type else 'Нет типа'}")
            document.add_paragraph(f"Исход: {mission.outcome if mission.outcome else 'Нет данных'}")  # Добавляем исход
            document.add_paragraph("-" * 40)

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = 'attachment; filename="missions.docx"'
        document.save(response)
        return response


    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

        popular_program = serializers.DictField(required=False)

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        # Основные статистические данные
        stats = Mission.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        # Получаем самую популярную космическую программу
        popular_program = (
            Mission.objects.values("space_program__name")
            .annotate(missions_count=Count("id"))
            .order_by("-missions_count")
            .first()
        )

        # Если данные о популярной программе существуют, добавляем их
        stats["popular_program"] = popular_program if popular_program else {"space_program__name": "Нет данных", "missions_count": 0}

        # Сериализуем данные
        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)

class SpaceProgramViewset(mixins.CreateModelMixin,
                          mixins.DestroyModelMixin, 
                          mixins.UpdateModelMixin, 
                          mixins.RetrieveModelMixin, 
                          mixins.ListModelMixin, 
                          GenericViewSet):
    queryset = SpaceProgram.objects.all()
    serializer_class = SpaceProgramSerializer

    def get_queryset(self):
        qs = super().get_queryset()

        # Если пользователь — суперпользователь, возвращаем все записи
        if self.request.user.is_superuser:
            # Фильтрация по user_id (если указано в запросе)
            user_id = self.request.query_params.get('user_id')
            if user_id:
                qs = qs.filter(user_id=user_id)

        # Если пользователь аутентифицирован, фильтруем по текущему пользователю
        elif self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)

        # Для анонимных пользователей возвращаем пустой queryset
        else:
            return qs.none()

        # Применение фильтров по дополнительным параметрам
        name = self.request.query_params.get('name')
        if name:
            qs = qs.filter(name__icontains=name)  # Частичное совпадение по названию

        start_date = self.request.query_params.get('start_date')
        if start_date:
            qs = qs.filter(start_date=start_date)  # Фильтрация по дате начала

        end_date = self.request.query_params.get('end_date')
        if end_date:
            qs = qs.filter(end_date=end_date)  # Фильтрация по дате окончания

        description = self.request.query_params.get('description')
        if description:
            qs = qs.filter(description__icontains=description)  # Частичное совпадение по описанию

        return qs
    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = SpaceProgram.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer = self.StatsSerializer(instance=stats)

        return Response(serializer.data)




# Вьюсет для типов миссий
class MissionTypeViewset(mixins.CreateModelMixin,
                         mixins.DestroyModelMixin,  
                         mixins.UpdateModelMixin, 
                         mixins.RetrieveModelMixin, 
                         mixins.ListModelMixin, 
                         GenericViewSet):
    queryset = MissionType.objects.all()
    serializer_class = MissionTypeSerializer

    def get_queryset(self):
        qs = super().get_queryset()

        # Если пользователь — суперпользователь, возвращаем все записи
        if self.request.user.is_superuser:
            # Фильтрация по user_id (если указано в запросе)
            user_id = self.request.query_params.get('user_id')
            if user_id:
                qs = qs.filter(user_id=user_id)

        # Если пользователь аутентифицирован, фильтруем по текущему пользователю
        elif self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)

        # Для анонимных пользователей возвращаем пустой queryset
        else:
            return qs.none()

        # Применение фильтров по дополнительным параметрам
        name = self.request.query_params.get('name')
        if name:
            qs = qs.filter(name__icontains=name)  # Частичное совпадение по названию

        return qs

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = MissionType.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer = self.StatsSerializer(instance=stats)

        return Response(serializer.data)


# Вьюсет для мест запуска
class LaunchSiteViewset(mixins.CreateModelMixin,
                        mixins.DestroyModelMixin,  
                        mixins.UpdateModelMixin, 
                        mixins.RetrieveModelMixin, 
                        mixins.ListModelMixin, 
                        GenericViewSet):
    queryset = LaunchSite.objects.all()
    serializer_class = LaunchSiteSerializer

    def get_queryset(self):
        qs = super().get_queryset()

        # Если пользователь — суперпользователь, возвращаем все записи
        if self.request.user.is_superuser:
            # Фильтрация по user_id (если указано в запросе)
            user_id = self.request.query_params.get('user_id')
            if user_id:
                qs = qs.filter(user_id=user_id)

        # Если пользователь аутентифицирован, фильтруем по текущему пользователю
        elif self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)

        # Для анонимных пользователей возвращаем пустой queryset
        else:
            return qs.none()

        # Применение фильтров по дополнительным параметрам
        name = self.request.query_params.get('name')
        if name:
            qs = qs.filter(name__icontains=name)  # Частичное совпадение по названию

        location = self.request.query_params.get('location')
        if location:
            qs = qs.filter(location__icontains=location)  # Частичное совпадение по местоположению

        return qs

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = LaunchSite.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        serializer = self.StatsSerializer(instance=stats)

        return Response(serializer.data)


# Вьюсет для космических аппаратов
from rest_framework import mixins
from rest_framework.viewsets import GenericViewSet
from .models import SpaceCraft
from .serializers import SpaceCraftSerializer

class SpaceCraftViewset(mixins.CreateModelMixin,
                        mixins.DestroyModelMixin,  
                        mixins.UpdateModelMixin, 
                        mixins.RetrieveModelMixin, 
                        mixins.ListModelMixin, 
                        GenericViewSet):
    queryset = SpaceCraft.objects.all()
    serializer_class = SpaceCraftSerializer

    def get_queryset(self):
        qs = super().get_queryset()

        # Если пользователь — суперпользователь, возвращаем все записи
        if self.request.user.is_superuser:
            # Фильтрация по user_id (если указано в запросе)
            user_id = self.request.query_params.get('user_id')
            if user_id:
                qs = qs.filter(user_id=user_id)

        # Если пользователь аутентифицирован, фильтруем по текущему пользователю
        elif self.request.user.is_authenticated:
            qs = qs.filter(user=self.request.user)

        # Для анонимных пользователей возвращаем пустой queryset
        else:
            return qs.none()

        # Применение фильтров по дополнительным параметрам
        name = self.request.query_params.get('name')
        if name:
            qs = qs.filter(name__icontains=name)  # Частичное совпадение по названию

        manufacturer = self.request.query_params.get('manufacturer')
        if manufacturer:
            qs = qs.filter(manufacturer__icontains=manufacturer)  # Частичное совпадение по производителю

        launch_site = self.request.query_params.get('launch_site')
        if launch_site:
            qs = qs.filter(launch_site__id=launch_site)  # Фильтрация по космодрому (по id)

        mission = self.request.query_params.get('mission')
        if mission:
            qs = qs.filter(mission__id=mission)  # Фильтрация по миссии (по id)

        return qs

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()
        most_popular_manufacturer = serializers.CharField()
        manufacturer_count = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = SpaceCraft.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        # Вычисляем самого популярного производителя космических аппаратов
        popular_manufacturer = SpaceCraft.objects.values("manufacturer").annotate(
            manufacturer_count=Count("manufacturer")
        ).order_by("-manufacturer_count").first()

        # Добавляем данные о самом популярном производителе в статистику
        if popular_manufacturer:
            stats.update({
                "most_popular_manufacturer": popular_manufacturer["manufacturer"],
                "manufacturer_count": popular_manufacturer["manufacturer_count"]
            })
        else:
            stats.update({
                "most_popular_manufacturer": "Нет данных",
                "manufacturer_count": 0
            })

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)
    
class CurrentUserView(APIView):
    def get(self, request):
        return Response({
            'username': request.user.username,
            'email': request.user.email,
            'is_superuser': request.user.is_superuser,
    })
class UserViewSet(viewsets.ViewSet):
    def list(self, request):
        users = User.objects.all()  # Получаем всех пользователей
        serializer = UserSerializer(users, many=True)
        return Response(serializer.data)
    




from django.contrib.auth import authenticate, login, logout
from django.core.cache import cache
from django.utils import timezone
from django.conf import settings
from django.utils.decorators import method_decorator
from rest_framework.decorators import action
from rest_framework.permissions import AllowAny, IsAuthenticated, BasePermission
from rest_framework.response import Response
from rest_framework import status, serializers
import pyotp  # Не забудьте установить pyotp

# # Сериализатор для OTP
# class OTPSerializer(serializers.Serializer):
#     key = serializers.CharField()

# # Разрешение для проверки успешной двойной аутентификации
# class OTPRequired(BasePermission):
#     def has_permission(self, request, view):
#         otp_timestamp = cache.get(f'otp_timestamp_{request.user.id}')
#         # Ограничиваем действие двойной аутентификации до 10 минут
#         if otp_timestamp and (timezone.now() - otp_timestamp).total_seconds() < 600:
#             return True
#         return False

from django.contrib.auth import authenticate, login, logout
from django.core.cache import cache
from django.utils import timezone
from django.conf import settings
from rest_framework.decorators import action
from rest_framework.permissions import AllowAny, IsAuthenticated, BasePermission
from rest_framework.response import Response
from rest_framework import status, serializers
import pyotp  # Не забудьте установить pyotp

from django.utils import timezone



# class User2ViewSet(GenericViewSet):
#     permission_classes = [AllowAny]

#     class OTPSerializer(serializers.Serializer):
#         key = serializers.CharField()

#     class OTPRequired(BasePermission):
#         def has_permission(self, request, view):
#             otp_timestamp = cache.get(f'otp_good_{request.user.id}')
#             return otp_timestamp is not None

#     @action(detail=False, methods=["POST"], url_path="login", permission_classes=[])
#     def use_login(self, request, *args, **kwargs):
#         """
#         Обычный вход в систему с использованием имени пользователя и пароля.
#         """
#         username = request.data.get("username")
#         password = request.data.get("password")

#         if not username or not password:
#             return Response({"success": False, "error": "Username and password are required."},
#                             status=status.HTTP_400_BAD_REQUEST)

#         user = authenticate(request, username=username, password=password)
#         if user:
#             login(request, user)
#             return Response({"success": True, "userId": user.id, "username": user.username})
#         else:
#             return Response({"success": False, "error": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

#     @action(detail=False, methods=["POST"], url_path="otp-login", permission_classes=[IsAuthenticated], serializer_class=OTPSerializer)
#     def otp_login(self, request, *args, **kwargs):
#         """
#         Подтверждение OTP-кода.
#         """
#         serializer = self.get_serializer(data=request.data)
#         serializer.is_valid(raise_exception=True)

#         otp_key = serializer.validated_data.get("key")
#         totp = pyotp.TOTP(request.user.userprofile.otp_key)

#         if totp.verify(otp_key):
#             cache.set(f'otp_good_{request.user.id}', True, timeout=600)  # 10 минут
#             return Response({"success": True, "message": "OTP verified successfully."})
#         else:
#             return Response({"success": False, "error": "Invalid OTP code."}, status=status.HTTP_400_BAD_REQUEST)

#     @action(detail=False, methods=["GET"], url_path="otp-status", permission_classes=[IsAuthenticated])
#     def get_otp_status(self, request, *args, **kwargs):
#         """
#         Проверка, активна ли OTP-аутентификация для пользователя.
#         """
#         otp_good = cache.get(f'otp_good_{request.user.id}')
#         return Response({"otp_authenticated": otp_good is not None})

#     @action(detail=False, methods=["POST"], url_path="logout", permission_classes=[IsAuthenticated])
#     def logout_user(self, request, *args, **kwargs):
#         """
#         Выход из системы.
#         """
#         logout(request)
#         cache.delete(f'otp_good_{request.user.id}')  # Удаление OTP-аутентификации
#         return Response({"success": True, "message": "Logged out successfully."})

#     @action(detail=False, methods=["GET"], url_path="protected", permission_classes=[IsAuthenticated, OTPRequired])
#     def protected_action(self, request, *args, **kwargs):
#         """
#         Доступ к защищённому ресурсу после OTP-аутентификации.
#         """
#         return Response({"success": True, "message": "Access granted to protected resource."})



class OTPSerializer(serializers.Serializer):
    key = serializers.CharField()

class OTPRequired(BasePermission):
    def has_permission(self, request, view):
        otp_timestamp = cache.get(f'otp_good_{request.user.id}')
        return otp_timestamp is not None

class User2ViewSet(GenericViewSet):
    permission_classes = [AllowAny]
    serializer_class = UserSerializer

    def get_serializer_class(self):
        if self.action == 'login_user' and 'otp_key' in self.request.data:
            return OTPSerializer
        return super().get_serializer_class()

    # Вход в систему с поддержкой OTP
    @action(methods=["POST"], detail=False, url_path="login")
    def login_user(self, request, *args, **kwargs):
        username = request.data.get("username")
        password = request.data.get("password")
        otp_key = request.data.get("otp_key")

        if not username or not password:
            return Response({"success": False, "error": "Username and password are required."}, status=status.HTTP_400_BAD_REQUEST)
        
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            # Проверяем, требуется ли OTP для пользователя
            if otp_key:
                try:
                    totp = pyotp.TOTP(user.userprofile.otp_key)  # Убедитесь, что у пользователя есть otp_key
                    if not totp.verify(otp_key):
                        return Response({"success": False, "error": "Invalid OTP code."}, status=status.HTTP_400_BAD_REQUEST)
                    # Устанавливаем флаг двойной аутентификации
                    cache.set(f'otp_good_{user.id}', timezone.now(), timeout=600)  # Действие на 10 минут
                except AttributeError:
                    return Response({"success": False, "error": "OTP key not set for user."}, status=status.HTTP_400_BAD_REQUEST)

            login(request, user)
            return Response({
                "success": True,
                "username": user.username,
                "userId": user.id,
                "otp_authenticated": bool(otp_key)  # Указывает, что OTP-аутентификация пройдена
            })
        else:
            return Response({"success": False, "error": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

    # Получение информации о пользователе с флагом двойной аутентификации
    @action(methods=["GET"], detail=False, url_path="info")
    def get_info(self, request, *args, **kwargs):
        data = {
            "is_authenticated": request.user.is_authenticated,
            "otp_authenticated": cache.get(f'otp_good_{request.user.id}') is not None
        }
        if request.user.is_authenticated:
            data.update({
                "username": request.user.username,
                "userId": request.user.id
            })
        return Response(data)

    # Выход из системы
    @action(methods=["POST"], detail=False, url_path="logout")
    def logout_user(self, request, *args, **kwargs):
        logout(request)
        cache.delete(f'otp_good_{request.user.id}')  # Сброс OTP-авторизации
        return Response({"success": True, "message": "Successfully logged out."}, status=status.HTTP_200_OK)

    # Метод, доступный только после двойной аутентификации
    @action(detail=False, methods=['GET'], url_path='protected', permission_classes=[IsAuthenticated, OTPRequired])
    def protected_action(self, request, *args, **kwargs):
        return Response({"success": True, "message": "Access granted to protected action."})







# class User2ViewSet(GenericViewSet):
#     permission_classes = [AllowAny]  # Разрешаем доступ ко всем действиям, кроме тех, которые требуют аутентификации
    
#     # Указываем сериализатор по умолчанию
#     serializer_class = UserSerializer

#     def get_serializer_class(self):
#         if self.action == 'login_user':
#             return UserSerializer
#         return super().get_serializer_class()

#     # Вход в систему
#     @action(methods=["POST"], detail=False, url_path="login")
#     def login_user(self, request, *args, **kwargs):
#         username = request.data.get("username")
#         password = request.data.get("password")

#         if not username or not password:
#             return Response({"success": False, "error": "Username and password are required."}, status=status.HTTP_400_BAD_REQUEST)
        
#         user = authenticate(request, username=username, password=password)
#         if user is not None:
#             login(request, user)
#             return Response({
#                 "success": True,
#                 "username": user.username,
#                 "userId": user.id
#             })
#         else:
#             return Response({"success": False, "error": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

#     # Получение информации о пользователе
#     @action(methods=["GET"], detail=False, url_path="info")
#     def get_info(self, request, *args, **kwargs):
#         data = {
#             "is_authenticated": request.user.is_authenticated
#         }
#         if request.user.is_authenticated:
#             data.update({
#                 "username": request.user.username,
#                 "userId": request.user.id
#             })
#         return Response(data)

#     # Выход из системы
#     @action(methods=["POST"], detail=False, url_path="logout")
#     def logout_user(self, request, *args, **kwargs):
#         # Выход пользователя (не требует ввода имени или пароля)
#         logout(request)
#         return Response({"success": True, "message": "Successfully logged out."}, status=status.HTTP_200_OK)















# class User2ViewSet(GenericViewSet):
#     @action(url_path="", methods=["GET"], detail=False)
#     def get_info(self, request, *args, **kwargs):
#         data = {
#             "is_authenticated": request.user.is_authenticated
#         }
#         if request.user.is_authenticated:
#             data.update({
#                 "username": request.user.username,
#                 "userId": request.user.id  # Привели к camelCase
#             })
#         return Response(data)



# class UserProfileViewSet(
#     GenericViewSet
# ):
#     permission_classes = [IsAuthenticated]

#     class OTPSerializer(serializers.Serializer):
#         key = serializers.CharField()

#     class OTPRequired(BasePermission):
#         def has_permission(self, request, view):
#             return bool(request.user and cache.get('otp_good', False))
        
#     @action(detail=False, url_path="check-login", methods=['GET'], permission_classes=[])
#     def get_check_login(self, request, *args, **kwargs):
#         return Response({
#             'is_authenticated': self.request.user.is_authenticated
#         })
    
#     @action(detail=False, url_path="login", methods=['GET'], permission_classes=[])
#     def use_login(self, request, *args, **kwargs):
#         user= authenticate(username='username', password='pass')
#         if user:
#             login(request, user)
#         return Response({
#             'is_authenticated': bool(user)
#         })

#     @action(detail=False, url_path='otp-login', methods=['POST'], serializer_class=OTPSerializer)
#     def otp_login(self, *args, **kwargs):
#         totp = pyotp.TOTP(self.request.user.userprofile.opt_key)
        
#         serializer = self.get_serializer(data=self.request.data)
#         serializer.is_valid(raise_exception=True)

#         success = False
#         if totp.now() == serializer.validated_data['key']:
#             cache.set('otp_good', True, 10)
#             success = True

#         return Response({
#             'success': success
#         })
    
#     @action(detail=False, url_path='otp-status')
#     def get_otp_status(self, *args, **kwargs):
#         otp_good = cache.get('otp_good', False)
#         return Response({
#             'otp_good': otp_good
#         })
    
#     @action(detail=False, url_path='otp-required', permission_classes=[OTPRequired])
#     def page_with_otp_required(self, *args, **kwargs):
#         return Response({
#             'success': True
#         })